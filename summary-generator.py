import fitz   #for reading '.pdf' document
import docx   #for reading ''.docx' document
import re
import pandas as pd
import nltk
import networkx as nx
import numpy as np

#below function when you want to upload a file in UI

class DataLayer:
    def __init__(self,path):
        """
        This is the initialization function of the dataLayer class.
        
        :param 
        :param 
        :param 
        """
        self.path = path
        self.json_dict = {}
            
    def read_PDF(self):
        self.json_dict['file_type'] = 'PDF'
        doc = fitz.Document(self.path)
        pages = doc.pageCount
        main_list = []
        for p in range(pages):
            pg = doc.loadPage(p)
            text = pg.getText('text')
            text_split = text.split('\n')
            main_list.append(text_split)
#            main_list.append(text_split)
        self.json_dict['text'] = main_list
        return self.json_dict
    
    def read_doc(self):
        self.json_dict['file_type'] = 'DOC'
        doc = docx.Document(self.path)
        para_list = []
        for p in doc.paragraphs:
            para_list.append(p.text.strip())
        para_list = [item for item in para_list if item != '']
        self.json_dict['text'] = para_list
        return self.json_dict
            
    def read_txt_HTML(self):
        with open (self.path,'r') as file_text:
            text = file_text.read()
        if re.search(r'\<html\>',text,re.IGNORECASE):
            self.json_dict['file_type'] = 'HTML'
        else:
            self.json_dict['file_type'] = 'TXT'
        self.json_dict['text'] = text
        return self.json_dict
    
    def read_excel_CSV(self):
        if self.path.lower().endswith('.csv'):
            self.json_dict['file_type'] = 'CSV'
            df = pd.read_csv(self.path)
        elif self.path.lower().endswith('xlsx'):
            self.json_dict['file_type'] = 'XLSX'
            df = pd.read_excel(self.path)
        elif self.path.lower().endswith('.xls'):
            self.json_dict['file_type'] = 'XLS'
            df = pd.read_excel(self.path)
        temp_dict = df.to_dict(orient='dict')
        self.json_dict['text'] = temp_dict
        return self.json_dict
        
    def get_text(self):
        file_name_with_ext = re.split(r'\\|\/',self.path)[-1]
        file_name = file_name_with_ext[:file_name_with_ext.rfind('.')]
        if self.path.lower().endswith('.pdf'):
            self.json_dict = self.read_PDF()
        if self.path.lower().endswith('.docx') or self.path.lower().endswith('.doc'):
            self.json_dict = self.read_doc()
        if self.path.lower().endswith('.html') or self.path.lower().endswith('.txt'):
            self.json_dict = self.read_txt_HTML()
        if (self.path.lower().endswith('.csv') or self.path.lower().endswith('.xlsx')
               or self.path.lower().endswith('.xls')):
           self.json_dict = self.read_excel_CSV() 
        return file_name,self.json_dict
    

data_obj = DataLayer(path)
file_name,some_dict = data_obj.get_text() 

#path=('...')



    
#===============================================================================
#===============================================================================
#                           common function
#===============================================================================
class Extractinglayer:
        
        def __init__(self,some_dict):
            self.some_dict=some_dict
            self.data=""
        
        def Pdfdoc(self):
            #list_values = [ v for v in some_dict['text'].values() ]   extracting lists from dict
            list_values=self.some_dict['text']
            result = sum(list_values,[])
            str_concat = " ".join(result)
            data=re.sub(' +',' ',str_concat)
            return data 
        
        def Docx(self):
            ini_list=self.some_dict['text']
            str_concat = " ".join(ini_list)
            data=re.sub(' +',' ',str_concat)
            return data
        
        def Txt(self):
             txt_data=self.some_dict['text']
             return txt_data
        
        def Main(self):
            if self.some_dict['file_type']=="PDF":
                self.data = self.Pdfdoc()
            if self.some_dict['file_type']=="DOC":
                self.data = self.Docx()
            if self.some_dict['file_type']=="TXT":
                self.data = self.Txt()
            return self.data
                

data_obj2 = Extractinglayer(some_dict)
article_text = data_obj2.Main()
#=================================================================================================
#INSTEAD DOING ALL THESE YOU CAN GIVE SOME TEXT IN article_text='...' variable 
#================================================================================================

from nltk.tokenize import sent_tokenize, word_tokenize
sentences =nltk.sent_tokenize(article_text)
no_of_sentences=len(sentences)

#FOR COMBINING ALL THE LIST OF SENT INTO SERIES
#sentences = [y for x in sentences for y in x] # flatten list
# Extract word vectors
word_embeddings = {}
f = open(r'C:\Users\bollud\Downloads\glove.6B\glove.6B.300d.txt',encoding="utf8")
for line in f:
    values = line.split()
    word = values[0]
    coefs = np.asarray(values[1:], dtype='float32')
    word_embeddings[word] = coefs
f.close()

len(word_embeddings)

# remove punctuations, numbers and special characters
clean_sentences = pd.Series(sentences).str.replace("[^a-zA-Z]", " ")

# make alphabets lowercase
clean_sentences = [s.lower() for s in clean_sentences]

#nltk.download('stopwords')
from nltk.corpus import stopwords
stop_words = stopwords.words('english')

#fuction to remove these stopwords
def remove_stopwords(sen):
    sen_new=" ".join([i for i in sen if i not in stop_words])
    return sen_new
    
    
# remove stopwords from the sentences
clean_sentences = [remove_stopwords(r.split()) for r in clean_sentences]

sentence_vectors = []
for i in clean_sentences:
  if len(i) != 0:
    v = sum([word_embeddings.get(w, np.zeros((300,))) for w in i.split()])/(len(i.split())+0.001)
  else:
    v = np.zeros((300,))
  sentence_vectors.append(v)
 
    
# similarity matrix
sim_mat = np.zeros([len(sentences), len(sentences)])
from sklearn.metrics.pairwise import cosine_similarity

for i in range(len(sentences)):
  for j in range(len(sentences)):
    if i != j:
      sim_mat[i][j] = cosine_similarity(sentence_vectors[i].reshape(1,300), sentence_vectors[j].reshape(1,300))[0,0]

#==============================================================================
#      applying page rank algorithm
#==============================================================================
      
import networkx as nx

nx_graph = nx.from_numpy_array(sim_mat)
scores = nx.pagerank(nx_graph)

ranked_sentences = sorted(((scores[i],s) for i,s in enumerate(sentences)), reverse=True)

# Extract top 10 sentences as the summary
print("The number of sentences in the document: ", no_of_sentences)
n=int(input("How many sentences of summary do you need:"))
print("\n")

for i in range(n):
    print(ranked_sentences[i][1])